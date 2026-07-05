"""
Generates a professionally formatted Credit Appraisal Memorandum as a .docx
file using python-docx. Returns a BytesIO buffer ready to stream as a download.
"""
from __future__ import annotations

import io
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

# ── Colour palette ──────────────────────────────────────────────────────────
C_NAVY   = RGBColor(0x1A, 0x2A, 0x4A)
C_TEAL   = RGBColor(0x0D, 0x7C, 0x7C)
C_LIGHT  = RGBColor(0xE8, 0xF4, 0xF4)
C_FAIL   = RGBColor(0xC0, 0x20, 0x20)
C_PASS   = RGBColor(0x1A, 0x7A, 0x40)
C_WARN   = RGBColor(0xB8, 0x6E, 0x00)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_HEADER = RGBColor(0x1A, 0x2A, 0x4A)
C_ALT    = RGBColor(0xF2, 0xF6, 0xFA)


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "C0C8D8")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _p(doc: Document, text: str, bold=False, size=10, color=None, align=None, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def _section_heading(doc: Document, title: str):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = C_NAVY
    # bottom border via paragraph border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A7A7A")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _kv_table(doc: Document, rows: List[Tuple[str, str]]):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    col_widths = [Inches(2.4), Inches(4.0)]
    for i, (label, value) in enumerate(rows):
        row = table.rows[i]
        row.height = Pt(16)
        # Key cell
        kc = row.cells[0]
        kc.width = col_widths[0]
        _set_cell_bg(kc, "EEF2F8")
        _set_cell_borders(kc)
        kp = kc.paragraphs[0]
        kr = kp.add_run(label)
        kr.bold = True
        kr.font.size = Pt(9)
        kr.font.color.rgb = C_NAVY
        kc.paragraphs[0].paragraph_format.space_before = Pt(2)
        kc.paragraphs[0].paragraph_format.space_after = Pt(2)
        # Value cell
        vc = row.cells[1]
        vc.width = col_widths[1]
        _set_cell_bg(vc, "FFFFFF")
        _set_cell_borders(vc)
        vp = vc.paragraphs[0]
        vr = vp.add_run(str(value))
        vr.font.size = Pt(9)
        vc.paragraphs[0].paragraph_format.space_before = Pt(2)
        vc.paragraphs[0].paragraph_format.space_after = Pt(2)


def _data_table(doc: Document, headers: List[str], rows: List[Tuple], col_widths_in: List[float]):
    total_rows = 1 + len(rows)
    table = doc.add_table(rows=total_rows, cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hrow = table.rows[0]
    for j, h in enumerate(headers):
        cell = hrow.cells[j]
        cell.width = Inches(col_widths_in[j])
        _set_cell_bg(cell, "1A2A4A")
        _set_cell_borders(cell)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = C_WHITE
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    # Data rows
    for i, row_data in enumerate(rows):
        row = table.rows[i + 1]
        bg = "F2F6FA" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.width = Inches(col_widths_in[j])
            _set_cell_bg(cell, bg)
            _set_cell_borders(cell)
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(8.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)


def generate_cam_docx(cam: Dict[str, Any], analyst_narrative: str) -> io.BytesIO:
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    meta = cam["meta"]

    # ── Cover / Title ─────────────────────────────────────────────────────
    _p(doc, "CREDIT APPRAISAL MEMORANDUM", bold=True, size=16,
       color=C_NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)
    _p(doc, "Working Capital Facility", bold=False, size=11,
       color=C_TEAL, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    _p(doc, meta["company_name"], bold=True, size=13,
       color=C_NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _p(doc, f"Case ID: {meta['case_id']}   |   Generated: {meta['generated_at']}",
       size=8.5, color=RGBColor(0x70, 0x80, 0x90),
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _p(doc, f"Prepared by: {meta['prepared_by']}",
       size=8, color=RGBColor(0x90, 0xA0, 0xB0),
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    doc.add_page_break()

    # ── S1 Basic Information ──────────────────────────────────────────────
    _section_heading(doc, cam["s1_basic_info"]["title"])
    _kv_table(doc, cam["s1_basic_info"]["rows"])

    # ── S2 Credit Requirement ─────────────────────────────────────────────
    _section_heading(doc, cam["s2_credit_requirement"]["title"])
    _kv_table(doc, cam["s2_credit_requirement"]["rows"])

    # ── S3 Financial Analysis ─────────────────────────────────────────────
    _section_heading(doc, cam["s3_financial_analysis"]["title"])
    s3 = cam["s3_financial_analysis"]
    _p(doc, f"Period: {s3['period']}   |   Audited: {s3['is_audited']}", size=8.5,
       color=RGBColor(0x50, 0x60, 0x70), space_after=3)
    _p(doc, "P&L Summary", bold=True, size=9, color=C_NAVY, space_before=2, space_after=2)
    _kv_table(doc, s3["pnl_rows"])
    _p(doc, "Balance Sheet Summary", bold=True, size=9, color=C_NAVY, space_before=4, space_after=2)
    _kv_table(doc, s3["bs_rows"])
    _p(doc, "Ledger Analysis", bold=True, size=9, color=C_NAVY, space_before=4, space_after=2)
    _kv_table(doc, s3["ledger_rows"])

    # ── S4 Banking ────────────────────────────────────────────────────────
    _section_heading(doc, cam["s4_banking"]["title"])
    _kv_table(doc, cam["s4_banking"]["rows"])

    # ── S5 Bureau ─────────────────────────────────────────────────────────
    _section_heading(doc, cam["s5_bureau"]["title"])
    _kv_table(doc, cam["s5_bureau"]["rows"])

    # ── S6 GST ────────────────────────────────────────────────────────────
    _section_heading(doc, cam["s6_gst"]["title"])
    _kv_table(doc, cam["s6_gst"]["rows"])

    # ── S7 Triangulation ──────────────────────────────────────────────────
    _section_heading(doc, cam["s7_triangulation"]["title"])
    if cam["s7_triangulation"]["pairwise_rows"]:
        _p(doc, "Pairwise Turnover Comparison", bold=True, size=9, color=C_NAVY, space_before=2, space_after=2)
        _data_table(doc,
            ["Comparison", "Source A", "Source B", "Variance", "Trust Weight"],
            cam["s7_triangulation"]["pairwise_rows"],
            [2.2, 1.3, 1.3, 0.9, 1.1])
    if cam["s7_triangulation"]["source_weights"]:
        _p(doc, "Aggregated Source Trust Weights", bold=True, size=9, color=C_NAVY, space_before=4, space_after=2)
        _data_table(doc, ["Source", "Trust Weight"],
            cam["s7_triangulation"]["source_weights"], [2.5, 2.0])
    _p(doc, "Effective Metrics", bold=True, size=9, color=C_NAVY, space_before=4, space_after=2)
    _kv_table(doc, cam["s7_triangulation"]["effective_rows"])

    # ── S8 Fraud ──────────────────────────────────────────────────────────
    s8 = cam["s8_fraud"]
    _section_heading(doc, s8["title"])
    fraud_color = {"LOW": C_PASS, "MEDIUM": C_WARN, "HIGH": C_FAIL}.get(s8["fraud_risk"], C_NAVY)
    _p(doc, f"Overall Fraud Risk: {s8['fraud_risk']}", bold=True, size=10, color=fraud_color, space_after=4)
    if s8["signals"]:
        _p(doc, "Fraud Signals Detected", bold=True, size=9, color=C_NAVY, space_before=2, space_after=2)
        _data_table(doc, ["Signal Type", "Severity", "Detail"],
            s8["signals"], [1.8, 0.9, 4.1])
    else:
        _p(doc, "No fraud signals detected.", size=9, color=C_PASS)
    if s8["contradictions"]:
        _p(doc, "Contradictions (Source Disagreements)", bold=True, size=9, color=C_NAVY, space_before=4, space_after=2)
        _data_table(doc, ["Pair", "Variance", "Detail"],
            s8["contradictions"], [1.8, 0.9, 4.1])

    # ── S9 Policy ─────────────────────────────────────────────────────────
    s9 = cam["s9_policy"]
    _section_heading(doc, s9["title"])
    dec_color = {"CLEAR": C_PASS, "DEVIATION REQUIRED": C_WARN, "POLICY REJECT": C_FAIL}.get(s9["decision"], C_NAVY)
    _p(doc, f"Policy Decision: {s9['decision']}   |   Pass Rate: {s9['pass_rate']}   |   Deviation Flag: {s9['deviation_flag']}",
       bold=True, size=10, color=dec_color, space_after=4)
    if s9["rule_rows"]:
        _data_table(doc, ["Rule", "Result", "Actual Value", "Threshold", "Note"],
            s9["rule_rows"], [2.4, 0.7, 1.1, 0.9, 1.7])

    # ── S10 ML ────────────────────────────────────────────────────────────
    _section_heading(doc, cam["s10_ml"]["title"])
    _kv_table(doc, cam["s10_ml"]["rows"])

    # ── S11 Limit Optimisation ────────────────────────────────────────────
    s11 = cam["s11_limit"]
    _section_heading(doc, s11["title"])
    if s11["constraint_rows"]:
        _data_table(doc, ["Constraint", "Ceiling Amount"],
            s11["constraint_rows"], [3.5, 2.0])
    _p(doc, "", space_after=4)
    _kv_table(doc, [
        ("Recommended Sanction Limit", s11["recommended_limit"]),
        ("Binding Constraint", s11["binding_constraint"]),
        ("Note", s11["note"]),
    ])

    # ── S12 Analyst Observations ──────────────────────────────────────────
    doc.add_page_break()
    _section_heading(doc, "SECTION 12 — ANALYST OBSERVATIONS")
    _p(doc, "The following narrative was generated by the Credit Decisioning OS AI pipeline "
       "based on verified data from the underwriting process. All figures referenced below "
       "are drawn from extracted and reconciled source data.",
       size=8.5, color=RGBColor(0x70, 0x80, 0x90), space_after=8)
    for para in (analyst_narrative or "").split("\n\n"):
        para = para.strip()
        if not para:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(para)
        r.font.size = Pt(9.5)

    # ── Footer note ───────────────────────────────────────────────────────
    doc.add_paragraph()
    _p(doc, "─" * 90, size=7, color=RGBColor(0xC0, 0xC8, 0xD8), space_before=8, space_after=2)
    _p(doc, f"CONFIDENTIAL — FOR INTERNAL CREDIT COMMITTEE USE ONLY   |   {meta['prepared_by']}",
       size=7.5, color=RGBColor(0x80, 0x90, 0xA0),
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
